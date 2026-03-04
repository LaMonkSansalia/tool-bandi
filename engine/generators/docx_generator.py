"""
DOCX Generator — generates Word documents using python-docx.

Generates documents programmatically (no binary template files needed).
Covers: proposta_tecnica, dichiarazione_sostitutiva.

Usage:
    from engine.generators.docx_generator import generate_docx

    docx_bytes = generate_docx("proposta_tecnica", context)
    with open("output.docx", "wb") as f:
        f.write(docx_bytes)
"""
from __future__ import annotations
import io
import logging
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)

# Colors — Italian PA formal blue
PA_BLUE = RGBColor(0x1A, 0x1A, 0x6E)
WARNING_YELLOW = RGBColor(0xFF, 0xC1, 0x07)
WARNING_TEXT = RGBColor(0x85, 0x64, 0x04)
PLACEHOLDER_BG = RGBColor(0xFF, 0xF3, 0xCD)

PLACEHOLDER_TEXT = "⚠ DA COMPILARE MANUALMENTE"


def _apply_heading_style(paragraph, level: int = 1) -> None:
    """Apply heading style with PA blue color."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.bold = True
    run.font.color.rgb = PA_BLUE
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)


def _add_table_row(table, label: str, value: str | None) -> None:
    """Add a two-column table row with bold label."""
    row = table.add_row()
    label_cell = row.cells[0]
    value_cell = row.cells[1]

    label_run = label_cell.paragraphs[0].add_run(label)
    label_run.bold = True
    label_run.font.color.rgb = PA_BLUE

    val = value or PLACEHOLDER_TEXT
    value_run = value_cell.paragraphs[0].add_run(val)
    if not value:
        value_run.font.color.rgb = WARNING_TEXT


def _add_placeholder(doc: Document, text: str = "") -> None:
    """Add a visually distinct placeholder paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(f"⚠ DA COMPILARE MANUALMENTE: {text}")
    run.font.color.rgb = WARNING_TEXT
    run.font.bold = True
    run.font.size = Pt(10)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)


def _get_value(obj: Any, *keys, default: str = "") -> str:
    """Safe nested dict/object value access."""
    current = obj
    for key in keys:
        if isinstance(current, dict):
            current = current.get(key)
        elif hasattr(current, key):
            current = getattr(current, key)
        else:
            return default
        if current is None:
            return default
    return str(current) if current is not None else default


# ─────────────────────────────────────────────
# PROPOSTA TECNICA
# ─────────────────────────────────────────────

def _build_proposta_tecnica(doc: Document, context: dict) -> None:
    """Build proposta tecnica document structure."""
    company = context.get("company", {})
    bando = context.get("bando", {})
    content = context.get("content", {})
    references = context.get("references", [])
    generated_at = context.get("generated_at", datetime.now().strftime("%d/%m/%Y"))
    version = context.get("version", "v1")

    # Title
    title = doc.add_heading("PROPOSTA TECNICA", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = PA_BLUE

    subtitle = doc.add_paragraph(_get_value(bando, "titolo", default=""))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    subtitle.runs[0].font.italic = True

    meta = doc.add_paragraph(
        f"Ente: {_get_value(bando, 'ente_erogatore')} | "
        f"Data: {generated_at} | Versione: {version} | "
        f"BOZZA - da revisionare prima dell'invio"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # 1. Soggetto proponente
    doc.add_heading("1. Soggetto Proponente", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(5.5)
    table.columns[1].width = Cm(10)

    rl = company.get("rappresentante_legale", {}) if isinstance(company, dict) else {}
    sede = company.get("sede", {}) if isinstance(company, dict) else {}

    rows_data = [
        ("Denominazione", _get_value(company, "denominazione")),
        ("Forma giuridica", _get_value(company, "forma_giuridica")),
        ("P.IVA", _get_value(company, "partita_iva")),
        ("Codice Fiscale", _get_value(company, "codice_fiscale")),
        ("ATECO", _get_value(company, "ateco_primario")),
        ("Sede legale", f"{_get_value(sede, 'indirizzo')}, {_get_value(sede, 'cap')} {_get_value(sede, 'comune')} ({_get_value(sede, 'provincia')})"),
        ("Rappresentante legale", f"{_get_value(rl, 'nome')} {_get_value(rl, 'cognome')}"),
        ("PEC", _get_value(company, "pec")),
    ]
    for label, value in rows_data:
        _add_table_row(table, label, value)

    doc.add_paragraph()

    # 2. Descrizione del progetto
    doc.add_heading("2. Descrizione del Progetto Proposto", level=1)
    desc = content.get("descrizione_progetto") if isinstance(content, dict) else None
    if desc:
        doc.add_paragraph(desc)
    else:
        _add_placeholder(doc, "Descrivere il progetto proposto in relazione agli obiettivi del bando.")

    # 3. Competenze e referenze
    doc.add_heading("3. Competenze e Referenze", level=1)
    doc.add_heading("3.1 Competenze Tecniche", level=2)
    comp_text = content.get("competenze_tecniche") if isinstance(content, dict) else None
    if comp_text:
        doc.add_paragraph(comp_text)
    else:
        _add_placeholder(doc, "Inserire le competenze tecniche rilevanti per questo bando.")

    doc.add_heading("3.2 Referenze Progettuali", level=2)
    if references:
        ref_table = doc.add_table(rows=1, cols=4)
        ref_table.style = "Table Grid"
        header_row = ref_table.rows[0]
        for i, hdr in enumerate(["Progetto", "Cliente/Settore", "Anno", "Descrizione"]):
            cell = header_row.cells[i]
            run = cell.paragraphs[0].add_run(hdr)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Note: python-docx doesn't easily set cell background in one line,
            # but the content is correct.

        for ref in references:
            row = ref_table.add_row()
            row.cells[0].paragraphs[0].add_run(_get_value(ref, "nome")).bold = True
            row.cells[1].paragraphs[0].add_run(_get_value(ref, "cliente_settore", default="—"))
            row.cells[2].paragraphs[0].add_run(_get_value(ref, "anno", default="—"))
            row.cells[3].paragraphs[0].add_run(_get_value(ref, "descrizione"))
    else:
        doc.add_paragraph("Nessuna referenza verificata disponibile.")

    doc.add_paragraph()

    # 4. Metodologia
    doc.add_heading("4. Metodologia e Piano di Lavoro", level=1)
    metodo = content.get("metodologia") if isinstance(content, dict) else None
    if metodo:
        doc.add_paragraph(metodo)
    else:
        _add_placeholder(doc, "Descrivere metodologia e fasi di lavoro.")

    # 5. Risultati attesi
    doc.add_heading("5. Risultati Attesi e Indicatori", level=1)
    risultati = content.get("risultati_attesi") if isinstance(content, dict) else None
    if risultati:
        doc.add_paragraph(risultati)
    else:
        _add_placeholder(doc, "Descrivere i risultati attesi e gli indicatori di misura.")

    # 6. Budget
    doc.add_heading("6. Budget Dettagliato", level=1)
    _add_placeholder(doc, "Inserire il budget dettagliato con voci di spesa, importi e % cofinanziamento.")

    # 7. Dichiarazioni
    doc.add_heading("7. Dichiarazioni del Rappresentante Legale", level=1)
    doc.add_paragraph(
        f"Il/La sottoscritto/a {_get_value(rl, 'nome')} {_get_value(rl, 'cognome')}, "
        f"rappresentante legale di {_get_value(company, 'denominazione')}, "
        "dichiara che le informazioni contenute nella presente proposta sono veritiere e complete."
    )

    # Signature
    doc.add_paragraph()
    sig_para = doc.add_paragraph(f"Palermo, {generated_at}")
    doc.add_paragraph()
    doc.add_paragraph("Il Rappresentante Legale")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph(f"{_get_value(rl, 'nome')} {_get_value(rl, 'cognome')}")


# ─────────────────────────────────────────────
# DICHIARAZIONE SOSTITUTIVA
# ─────────────────────────────────────────────

def _build_dichiarazione(doc: Document, context: dict) -> None:
    """Build dichiarazione sostitutiva document."""
    company = context.get("company", {})
    bando = context.get("bando", {})
    generated_at = context.get("generated_at", datetime.now().strftime("%d/%m/%Y"))
    version = context.get("version", "v1")

    rl = company.get("rappresentante_legale", {}) if isinstance(company, dict) else {}
    sede = company.get("sede", {}) if isinstance(company, dict) else {}

    title = doc.add_heading("DICHIARAZIONE SOSTITUTIVA DI ATTO NOTORIO", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = PA_BLUE

    sub = doc.add_paragraph("ai sensi degli artt. 46 e 47 del D.P.R. 28 dicembre 2000 n. 445")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.italic = True

    doc.add_paragraph(
        f"Bando: {_get_value(bando, 'titolo')} | "
        f"Data: {generated_at} | Versione: {version}"
    ).runs[0].font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("Il/La sottoscritto/a:")

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("Cognome e nome", f"{_get_value(rl, 'cognome')} {_get_value(rl, 'nome')}"),
        ("Luogo e data di nascita", f"{_get_value(rl, 'luogo_nascita')} del {_get_value(rl, 'data_nascita')}"),
        ("Codice Fiscale", _get_value(rl, "codice_fiscale")),
        ("Residente a", _get_value(rl, "residenza")),
        ("In qualità di", f"Rappresentante Legale di {_get_value(company, 'denominazione')}"),
    ]
    for label, value in rows_data:
        _add_table_row(table, label, value or None)

    doc.add_paragraph()
    decl_heading = doc.add_paragraph("DICHIARA")
    decl_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    decl_heading.runs[0].bold = True
    decl_heading.runs[0].font.size = Pt(14)

    declarations = [
        "di non aver riportato condanne penali passate in giudicato;",
        "di non essere destinatario di misure di prevenzione iscritte nel casellario giudiziale;",
        f"che l'impresa {_get_value(company, 'denominazione')}, P.IVA {_get_value(company, 'partita_iva')}, è regolarmente iscritta alla CCIAA;",
        "che l'impresa è in regola con gli obblighi tributari e previdenziali (DURC regolare);",
        "che l'impresa non si trova in stato di fallimento o procedure concorsuali.",
    ]
    for i, decl in enumerate(declarations, 1):
        doc.add_paragraph(f"{i}. {decl}", style="List Number")

    doc.add_paragraph()
    doc.add_paragraph(
        f"Palermo, {generated_at}\n\n"
        "Il dichiarante\n"
        "_" * 40 + "\n"
        f"{_get_value(rl, 'nome')} {_get_value(rl, 'cognome')}\n\n"
        "Allegare copia documento di identità in corso di validità (art. 38 D.P.R. 445/2000)"
    ).runs[0].font.size = Pt(10)


# ─────────────────────────────────────────────
# PUBLIC API
# ─────────────────────────────────────────────

BUILDERS = {
    "proposta_tecnica": _build_proposta_tecnica,
    "dichiarazione_sostitutiva": _build_dichiarazione,
}


def generate_docx(template_name: str, context: dict) -> bytes:
    """
    Generate a Word document for the given template.

    Args:
        template_name: One of "proposta_tecnica", "dichiarazione_sostitutiva"
        context: Template variables

    Returns:
        DOCX as bytes

    Raises:
        ValueError: Unknown template
        RuntimeError: Generation error
    """
    if template_name not in BUILDERS:
        raise ValueError(
            f"Unknown DOCX template: '{template_name}'. "
            f"Available: {list(BUILDERS.keys())}"
        )

    # Inject common context
    full_context = {
        "generated_at": datetime.now().strftime("%d/%m/%Y"),
        **context,
    }

    doc = Document()

    # Page margins (A4 formal)
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    try:
        BUILDERS[template_name](doc, full_context)
    except Exception as e:
        raise RuntimeError(f"DOCX generation failed for '{template_name}': {e}") from e

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    logger.info(f"Generated DOCX '{template_name}': {len(docx_bytes)} bytes")
    return docx_bytes
